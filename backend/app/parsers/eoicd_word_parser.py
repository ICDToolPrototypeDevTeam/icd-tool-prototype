"""
EoICD Word 主文件解析器。

支持 .doc（通过 LibreOffice 转换为 .docx）和 .docx 格式。
提取段落（保留 Heading 层级）、内嵌表格、图片引用和接口信号信息。
"""

import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from app.models import EoICDChunk, ParsedEoICDInterface

# LibreOffice 常见安装路径
_LIBREOFFICE_PATHS = [
    "soffice",
    "soffice.exe",
    "libreoffice",
    "/usr/bin/soffice",
    "C:/Program Files/LibreOffice/program/soffice.exe",
    "C:/Program Files (x86)/LibreOffice/program/soffice.exe",
]


def _find_soffice() -> Optional[str]:
    """查找 LibreOffice soffice 可执行文件路径。"""
    for p in _LIBREOFFICE_PATHS:
        if shutil.which(p) or os.path.isfile(p):
            return p
    return None


def _is_doc_file(filepath: Path) -> bool:
    """判断是否为旧版 .doc 二进制格式（OLE Compound Document）。"""
    if filepath.suffix.lower() == ".doc":
        return True
    try:
        with open(filepath, "rb") as f:
            magic = f.read(8)
        return magic[:4] == b"\xd0\xcf\x11\xe0"
    except Exception:
        return False


def _convert_doc_to_docx(doc_path: Path, output_dir: Path) -> Path:
    """使用 LibreOffice 将 .doc 转换为 .docx。

    Args:
        doc_path: 输入 .doc 文件路径
        output_dir: 输出目录

    Returns:
        转换后的 .docx 文件路径

    Raises:
        RuntimeError: 若 LibreOffice 未安装或转换失败
    """
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError(
            "LibreOffice 未安装，无法解析 .doc 文件。"
            "请从 https://www.libreoffice.org/ 安装 LibreOffice。"
        )

    cmd = [
        soffice,
        "--headless",
        "--convert-to", "docx",
        "--outdir", str(output_dir),
        str(doc_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice .doc → .docx 转换失败 (exit={result.returncode}): "
            f"{result.stderr[:500]}"
        )

    for f in output_dir.iterdir():
        if f.suffix.lower() == ".docx":
            return f

    raise RuntimeError(
        f"LibreOffice 转换完成但未找到输出 .docx 文件，"
        f"输出目录内容: {list(output_dir.iterdir())}"
    )


def _resolve_docx_path(word_path: Path) -> Path:
    """将输入的 Word 文件解析为可读的 .docx 路径。

    - 若已是 .docx → 直接返回
    - 若是 .doc → 转换为临时 .docx 再返回
    """
    if not _is_doc_file(word_path):
        return word_path

    temp_dir = Path(tempfile.mkdtemp(prefix="eoicd_convert_"))
    return _convert_doc_to_docx(word_path, temp_dir)


def _style_to_heading_level(style_name: str, style_id: str) -> int:
    """从 python-docx 段落样式中提取 Heading 层级 (1-9)，非标题返回 0。"""
    s = style_name or ""
    s_lower = s.lower()

    # 标准 "Heading N" 样式
    if s_lower.startswith("heading "):
        try:
            return int(s.split()[-1])
        except ValueError:
            pass

    # .doc 转换后的中文自定义标题样式：！一级条 / ！二级条 / ……
    # 中文数字映射
    cn_num = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
    for cn, n in cn_num.items():
        if f"{cn}级条" in s or f"{cn}級條" in s:
            return n

    # TOC 标题样式视为 Heading 1-2
    if "toc" in s_lower:
        try:
            return int(s_lower.split()[-1])
        except ValueError:
            return 1

    # 其它包含 heading 关键字
    if "heading" in s_lower:
        try:
            return int(s_lower.split()[-1])
        except ValueError:
            pass

    # Outline level（段落属性中的大纲级别）
    if style_id:
        sid = style_id.lower()
        if "outlinelevel" in sid or "heading" in sid:
            try:
                return int(sid.split("outlinelevel")[-1].split()[0])
            except ValueError:
                pass

    return 0


def _format_paragraph_as_markdown(paragraph, heading_level: int) -> str:
    """将单个段落格式化为 Markdown 文本。"""
    text = paragraph.text.strip()
    if not text:
        return ""

    if heading_level > 0:
        prefix = "#" * min(heading_level, 6)
        return f"{prefix} {text}\n"
    else:
        return f"{text}\n\n"


def _extract_table_as_dict(table) -> dict:
    """将 python-docx Table 提取为结构化 dict。"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(c for c in cells):
            rows.append(cells)
    return {"rows": rows, "row_count": len(rows), "col_count": max((len(r) for r in rows), default=0)}


def _extract_interfaces_from_tables(tables: list[dict], source_file: str) -> list[ParsedEoICDInterface]:
    """从表格中识别信号/接口信息并转换为 ParsedEoICDInterface 列表。

    通过表头关键词（接口名/信号/方向/数据类型/周期等）识别接口表格。
    """
    interfaces: list[ParsedEoICDInterface] = []
    signal_keywords = {"信号", "signal", "接口", "interface", "pin", "针脚", "针孔", "参数", "parameter"}

    for table in tables:
        rows = table.get("rows", [])
        if len(rows) < 2:
            continue

        header = rows[0]
        header_lower = [h.lower() for h in header]

        # 检查是否为接口表格
        if not any(kw in " ".join(header_lower) for kw in signal_keywords):
            continue

        # 定位关键列
        col_map: dict[str, int] = {}
        col_aliases = {
            "interface_name": ["接口名", "interface name", "接口名称", "interface"],
            "interface_direction": ["方向", "direction", "dir", "发送/接收"],
            "signal_name": ["信号名", "signal name", "信号名称", "信号", "名称", "name"],
            "data_type": ["数据类型", "data type", "类型", "type"],
            "transfer_cycle": ["周期", "cycle", "刷新率", "refresh", "period", "传输周期"],
            "description": ["描述", "description", "说明", "备注", "功能"],
        }
        for field, aliases in col_aliases.items():
            for i, h in enumerate(header):
                if any(a in h.lower() or a in h for a in aliases):
                    col_map[field] = i
                    break

        if "signal_name" not in col_map:
            continue

        for row in rows[1:]:
            if not row or not any(c for c in row):
                continue
            try:
                interface = ParsedEoICDInterface(
                    interface_name=row[col_map.get("interface_name", 0)] if "interface_name" in col_map else "",
                    interface_direction=row[col_map["interface_direction"]] if "interface_direction" in col_map else "",
                    signal_name=row[col_map["signal_name"]],
                    data_type=row[col_map["data_type"]] if "data_type" in col_map else "",
                    transfer_cycle=row[col_map["transfer_cycle"]] if "transfer_cycle" in col_map else None,
                    source_file=source_file,
                    description=row[col_map["description"]] if "description" in col_map else None,
                )
                interfaces.append(interface)
            except IndexError:
                continue

    return interfaces


def _extract_images_from_docx(doc: DocxDocument) -> list[dict]:
    """提取文档中的图片引用信息。"""
    images = []
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            images.append({
                "index": i,
                "filename": rel.target_ref.split("/")[-1] if rel.target_ref else f"image_{i}",
                "content_type": rel.target_part.content_type if rel.target_part else "",
            })
    return images


def parse_eoicd_word(word_path: Path) -> list[EoICDChunk]:
    """解析 EoICD Word 主文件，返回 EoICDChunk 列表。

    当前版本将整个 Word 文件封装为 1 个 chunk-001。
    提取段落（Markdown）、表格（结构化 dict）、接口信息、图片引用。

    Args:
        word_path: EoICD Word 文件路径（支持 .doc 和 .docx）

    Returns:
        EoICDChunk 列表（当前长度为 1）
    """
    if not word_path.exists():
        raise FileNotFoundError(f"Word 文件不存在: {word_path}")

    # 1. 确保可读的 .docx
    docx_path = _resolve_docx_path(word_path)
    doc = DocxDocument(str(docx_path))

    # 2. 遍历段落和表格，构建 content 和 tables
    content_parts: list[str] = []
    tables: list[dict] = []
    heading_count = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # 段落
            para = _find_paragraph_by_xml(doc, element)
            if para is None:
                continue
            style_name = para.style.name if para.style else ""
            style_id = ""
            try:
                pPr = element.find(qn("w:pPr"))
                if pPr is not None:
                    pStyle = pPr.find(qn("w:pStyle"))
                    if pStyle is not None:
                        style_id = pStyle.get(qn("w:val"), "")
            except Exception:
                pass
            heading_level = _style_to_heading_level(style_name, style_id)
            text = _format_paragraph_as_markdown(para, heading_level)
            if text:
                content_parts.append(text)
                if heading_level > 0:
                    heading_count += 1

        elif tag == "tbl":
            # 表格
            table = _find_table_by_xml(doc, element)
            if table is not None:
                table_dict = _extract_table_as_dict(table)
                if table_dict["rows"]:
                    tables.append(table_dict)
                    # 表格也以 Markdown 表格形式写入 content
                    table_md = _table_to_markdown(table_dict)
                    content_parts.append(table_md)

    content = "".join(content_parts)

    # 3. 提取接口信息
    interfaces = _extract_interfaces_from_tables(tables, str(word_path))

    # 4. 提取图片引用
    images = _extract_images_from_docx(doc)

    # 5. 构建 chunk
    file_stem = word_path.stem
    context_summary = (
        f"本文档为 EoICD 接口控制文件（{file_stem}），"
        f"包含 {heading_count} 个标题段落、"
        f"{len(tables)} 个表格、{len(interfaces)} 个接口信号、{len(images)} 张图片。"
    )

    chunk = EoICDChunk(
        chunk_id="chunk-001",
        chunk_title=file_stem,
        source_file=str(word_path),
        source_section="全篇",
        source_page_range="全文",
        content=content,
        tables=tables,
        interfaces=interfaces,
        context_summary=context_summary,
    )

    return [chunk]


def _find_paragraph_by_xml(doc: DocxDocument, xml_element):
    """根据 XML 元素查找对应的 python-docx Paragraph 对象。"""
    for para in doc.paragraphs:
        if para._element is xml_element:
            return para
    return None


def _find_table_by_xml(doc: DocxDocument, xml_element):
    """根据 XML 元素查找对应的 python-docx Table 对象。"""
    for table in doc.tables:
        if table._element is xml_element:
            return table
    return None


def _table_to_markdown(table: dict) -> str:
    """将表格 dict 转为 Markdown 表格文本。"""
    rows = table.get("rows", [])
    if not rows:
        return ""

    max_cols = max((len(r) for r in rows), default=0)
    if max_cols == 0:
        return ""

    # 确保每行列数一致
    padded = [r + [""] * (max_cols - len(r)) for r in rows]

    lines = []
    # 表头
    lines.append("| " + " | ".join(padded[0]) + " |")
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    # 数据行
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n" + "\n".join(lines) + "\n\n"
