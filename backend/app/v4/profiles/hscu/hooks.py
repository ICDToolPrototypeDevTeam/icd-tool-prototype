"""HSCU-specific profile hooks.

Implements ``preprocess_hlr_requirements`` invoked by
``app.v4.profiles.apply_hlr_preprocess_hook`` after HLR Word parsing.

HSCU HLR text references ARINC-429 labels symbolically
(``LBL_DIS_00_SYS1``), but the corresponding EoICD blocks encode them with
the octal label number (``L145_DIS_00_SYS1_T1A``). The reverse matcher's
Stage1 label-prefix filter therefore cannot match HSCU requirements against
EoICD blocks. This hook appends an ``(alias: L<octal>_<NAME>)`` annotation
to each affected requirement's ``content`` so the AI labeler sees both forms
and emits the octal-form labels that the matcher recognises.

Two strategies are supported (declared in ``config.yaml`` under
``hlr_preprocess``):

1. ``extra_mappings`` — explicit dict of ``{LBL_XXX: octal}`` pairs
   maintained by hand.
2. ``auto_parse_hlr_table_0`` — auto-extract ``(LBL_XXX, octal)`` pairs
   from a "标签总览" / "LABEL catalog" table at the configured table
   index inside the HLR Word doc. The hook re-opens ``hlr_out.source_file``
   to read the docx tables; ``HLROutput`` itself only carries parsed
   requirements, not raw tables.

If both are configured, Table 0 mappings are used as the base and
``extra_mappings`` extend/override them.

Original text is preserved verbatim. The annotation is a single-line
append so existing field structure and downstream tokenisation are
unaffected.
"""
from __future__ import annotations

import re
from pathlib import Path

from app.v4.profiles.base import ControllerProfile


_LBL_TOKEN = re.compile(r"(?<![A-Z0-9_])LBL_([A-Z0-9_]+?)(?:_SSM)?(?![A-Z0-9_])")
# Bare signal token (no ``LBL_`` prefix). Matches RDCU1 catalog col-5
# signal names that appear in HSCU HLR body text without the LBL_ prefix
# (e.g. ``ABV1_LOAD_VOLT_AVAIL_RPDU_R1``, ``AIR_SPEED_FCM1_R1``).
# ``(?<!LBL_)`` rejects matches that are already captured as part of a
# ``LBL_<NAME>`` token (those go through the _LBL_TOKEN path instead).
_BARE_SIGNAL_TOKEN = re.compile(r"(?<!LBL_)(?<![A-Z0-9_])([A-Z][A-Z0-9_]{4,})(?![A-Z0-9_])")
_R_SUFFIX_RE = re.compile(r"_R\d+$")
_OCTAL_CELL_RE = re.compile(r"^L?(\d+)$")  # matches "145" or "L145"

# 8-column RDCU1 catalog (Table[0] in the new HSCU HLR Word) packs
# multiple signal names into one cell (col 5), one signal per paragraph.
# Each signal name is a bare token (no LBL_ prefix in HLR body text), so
# the bare-name path of the regex below wouldn't match. We extract them
# as separate mapping entries so HSCU HLR text that references a signal
# by its bare name (e.g. ``ABV1_LOAD_VOLT_AVAIL_RPDU_R1``) is aliased to
# ``L51_ABV1_LOAD_VOLT_AVAIL_RPDU_R1``.
#
# The signal-name heuristic rejects anything that isn't a pure
# uppercase/digit/underscore token of reasonable length. This excludes
# accidental matches against Chinese column headers, SSM names that
# happen to land in col 5, etc.
_SIGNAL_LINE_RE = re.compile(r"^[A-Z][A-Z0-9_]{4,}$")
# Index of the signal-name column in an 8-col RDCU1 catalog table. Verified
# against the new HSCU HLR Word: col1=LBL name, col2=OCT, col3=SDI,
# col4=通道, col5=信号名称 (one per line), col6=rate, col7=SSM. 4-col
# HSCU catalogs don't have a signal-name column.
_SIGNAL_COL_INDEX = 5
_MIN_SIGNAL_COLS = _SIGNAL_COL_INDEX + 1  # 6 cells needed to access col 5

# 目录表 SDI/SSM 列位置（两种表格式不同，按行单元格数区分）：
#   - 8 列表（RDCU1 inbound，Table[0]）：col3=SDI号, col7=SSM类型
#   - 4 列表（HSCU outbound，Table[8]）：col2=SDI号，无 SSM 列
_SDI_COL_8 = 3
_SSM_COL_8 = 7
_SDI_COL_4 = 2
# SDI 单值形式（0-3 单数字）才写入别名供匹配层值级比对；
# "1=HSCU_A;2=HSCU_B" 这类多值映射不写（无法归约为单值）。
_SDI_SINGLE_RE = re.compile(r"^[0-3]$")
_SSM_TOKEN_RE = re.compile(r"^[A-Za-z]{2,8}$")


def _bare_name(raw: str) -> str:
    """Strip a trailing ``_SSM`` suffix (ARINC-429 SSM bit) from a captured
    label name. HSCU HLR text writes ``LBL_DIS_00_SYS1_SSM`` (label + SSM
    bit field); the mapping key is the bare label ``LBL_DIS_00_SYS1``.
    """
    return raw[:-4] if raw.upper().endswith("_SSM") else raw


def _format_alias_octal(octal_value: str) -> str:
    """Return the octal as a left-padded 3-digit string.

    ARINC-429 label numbers are 8-bit values rendered as 3-digit octal
    (range ``000..377``). EoICD block keys use the 3-digit form
    (``L051_ABV1_...``); HSCU's HLR catalog sometimes drops the leading
    zero (``51`` instead of ``051``). The Stage1 label-prefix filter in
    ``reverse_matcher`` compares ``L<label>/`` against block keys, so the
    alias must match EoICD's leading-zero form or the filter rejects the
    block. Padding once at alias-format time is cheaper and clearer than
    padding the mapping values everywhere they're used.
    """
    return octal_value.zfill(3)


def _looks_like_label_cell(text: str) -> bool:
    """True iff the cell text starts with ``LBL_`` followed by at least one
    alphanumeric char. Header cells ("LBL名称") do not match.
    """
    t = text.strip()
    if len(t) < 5:
        return False
    return t[:4].upper() == "LBL_" and (t[4:5].isalnum() or t[4:5] == "_")


def _looks_like_octal_cell(text: str) -> bool:
    """True iff the cell text is an octal-like ARINC-429 label number.

    Must satisfy ALL of:
      - Pure decimal digits or ``L`` prefix + digits (e.g. ``145`` / ``L145``)
      - No ``?`` (placeholders like ``???``)
      - Length >= 2 digits. ARINC-429 label numbers are 8-bit values
        rendered as octal (3 digits, range 000..377, or 1-2 digits when
        the leading zero is dropped — e.g. ``74`` is ``074``). 1-digit
        values are typically SDI codes (``0`` / ``1`` / ``2`` / ``3``)
        or refresh-rate fields, not label octal numbers.

    This length floor prevents SDI cells (typically single-digit
    ``0``/``1``/``2``/``3``) from being misread as label numbers when
    both columns are scanned within a single row.
    """
    t = text.strip()
    if not t or "?" in t:
        return False
    m = _OCTAL_CELL_RE.match(t)
    if not m:
        return False
    return len(m.group(1)) >= 2


def _identify_label_tables(tables: list) -> list[int]:
    """Return indices of ALL LBL catalog tables, in document order.

    A document may legitimately contain MULTIPLE LBL catalogs of different
    formats (e.g. HSCU HLR has Table[0] = RDCU1 inbound catalog with
    ``_R1`` suffix in 11x8 layout, AND Table[8] = HSCU outbound catalog
    without ``_R1`` suffix in 13x4 layout). Both are valid "label overview"
    tables and must be merged.

    Heuristic — a table qualifies iff:
      1. **Wide table**: column count >= 3 (excludes 8x2 requirement tables).
      2. **Multi-row**: row count >= 2 (header + at least one data row).
      3. **LBL data rows**: at least one row contains an LBL_* cell AND an
         octal-like numeric cell.

    Returns an empty list when no candidate table satisfies the heuristic.
    Document order is preserved so callers can deterministically resolve
    collisions (later entries naturally override earlier ones).
    """
    indices: list[int] = []
    for idx, tbl in enumerate(tables):
        # Heuristic 1: wide table (>= 3 cols excludes 8x2 requirement tables)
        if len(tbl.columns) < 3:
            continue
        # Heuristic 2: multi-row
        if len(tbl.rows) < 2:
            continue

        # Score: count rows that contain both an LBL_* cell and an octal cell
        lbl_rows = 0
        for row in tbl.rows:
            cell_texts = [c.text.strip() for c in row.cells]
            has_lbl = any(_looks_like_label_cell(t) for t in cell_texts)
            has_octal = any(_looks_like_octal_cell(t) for t in cell_texts)
            if has_lbl and has_octal:
                lbl_rows += 1

        if lbl_rows > 0:
            indices.append(idx)
    return indices


def _extract_row_mapping(cells: list[str]) -> tuple[str, str] | None:
    """From a single row's cell texts, extract ``(bare_name, octal_value)``.

    Uses within-row scanning (NOT fixed column indices):
      - First cell starting with ``LBL_`` → bare_name source
      - First cell matching ``<digits>`` or ``L<digits>`` → octal value

    Returns None if no valid (LBL, octal) pair is found in the row.
    Placeholder octal values (``???`` etc.) are filtered upstream.
    """
    bare_name: str | None = None
    octal_value: str | None = None

    for cell in cells:
        t = cell.strip()
        if bare_name is None and _looks_like_label_cell(t):
            bare_name = t
        if octal_value is None and _looks_like_octal_cell(t):
            m = _OCTAL_CELL_RE.match(t)
            if m:
                octal_value = m.group(1)
        if bare_name is not None and octal_value is not None:
            break

    if bare_name is None or octal_value is None:
        return None
    return bare_name, octal_value


def _extract_signal_names(cell_text: str) -> list[str]:
    """Parse a multi-line signal-name cell (8-col RDCU1 catalog col 5).

    The RDCU1 inbound catalog packs every signal carried by a single LBL
    into one cell, one signal name per paragraph. Each signal name is a
    bare token (no ``LBL_`` prefix) and HSCU HLR body text frequently
    references the signal by its bare name (e.g. ``ABV1_LOAD_VOLT_AVAIL_RPDU_R1``
    instead of ``LBL_ABV1_RPDU_R1``). Mapping these bare signal names to
    the row's octal lets the hook alias HSCU HLR text that uses the bare
    form back to the EoICD ``L<octal>_<signal>`` block key.

    Only tokens that match ``_SIGNAL_LINE_RE`` are returned. This rejects
    accidental matches against Chinese column headers, rate-template strings
    (``LABELRATE_200_TIMEOUT`` is a value, not a signal name), or stray
    punctuation that could land in col 5 in some documents.
    """
    out: list[str] = []
    for line in cell_text.splitlines():
        t = line.strip()
        if not t:
            continue
        if not _SIGNAL_LINE_RE.match(t):
            continue
        out.append(t)
    return out


def _extract_label_meta(
    source_file: str,
    table_index: int | None = None,
) -> dict[str, dict]:
    """Open ``source_file`` (HLR Word), auto-detect ALL LBL catalog tables,
    and return a ``{name: {"octal": str, "sdi": str, "ssm": str}}`` dict.

    Catalog tables are identified by ``_identify_label_tables``:
      - >= 3 columns (excludes 8x2 requirement tables)
      - >= 2 rows (header + at least one data row)
      - at least one row contains both an LBL_* cell and an octal cell

    A document may legitimately contain MULTIPLE catalogs of different
    formats with DIFFERENT column layouts (HSCU HLR: Table[0]=8-col
    RDCU1 inbound catalog with SDI号/SSM类型/信号名称 columns;
    Table[8]=4-col HSCU outbound catalog with SDI号 only). SDI/SSM
    columns are therefore resolved per-format by row cell count:
      - >= 8 cells → SDI at index 3, SSM at index 7 (Table[0] format)
      - >= 4 cells → SDI at index 2, no SSM (Table[8] format)
    Only single-digit SDI values (0-3) are kept; multi-value mappings
    like ``1=HSCU_A;2=HSCU_B`` are dropped (not reducible to one value).
    SSM is kept when it is a short alphabetic token (BNR/DIS/...).

    For each qualifying row, ``_extract_row_mapping`` finds the (LBL_*,
    octal) pair by within-row scanning (no fixed column assumption).
    Tables are processed in document order; on name collision, later
    tables override earlier ones field-by-field (fields present in the
    later row win; missing fields keep the earlier value).

    Both the original ``LBL_<NAME>`` form and the ``_R1``-stripped form
    are inserted, so downstream HLR text matching works whether the body
    writes ``LBL_<NAME>_R1`` (catalog form) or ``LBL_<NAME>`` (body form).
    Signal names from the 8-col format's col 5 inherit their row's
    octal/SDI/SSM.

    ``table_index`` is an optional escape hatch: when provided, skip the
    auto-detect heuristic and use ONLY the table at that index (caller
    asserts the index is known-good). Default ``None`` (auto, all
    matching tables).

    Returns an empty dict if the docx is unreadable, has no qualifying
    table, or all rows lack valid (LBL, octal) pairs.
    """
    result: dict[str, dict] = {}
    p = Path(source_file)
    if not p.exists():
        return result

    try:
        from docx import Document  # local import to keep hook import-time cheap
        d = Document(str(p))
    except Exception:
        return result

    chosen_indices: list[int]
    if table_index is not None:
        if table_index < 0 or table_index >= len(d.tables):
            return result
        chosen_indices = [table_index]
    else:
        chosen_indices = _identify_label_tables(d.tables)
    if not chosen_indices:
        return result

    for idx in chosen_indices:
        tbl = d.tables[idx]
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            pair = _extract_row_mapping(cells)
            if pair is None:
                continue
            raw_name, octal_value = pair

            # SDI/SSM columns resolved per table format (see docstring).
            sdi_value = ""
            ssm_value = ""
            if len(cells) >= 8:
                if _SDI_SINGLE_RE.match(cells[_SDI_COL_8]):
                    sdi_value = cells[_SDI_COL_8]
                if _SSM_TOKEN_RE.match(cells[_SSM_COL_8]):
                    ssm_value = cells[_SSM_COL_8].upper()
            elif len(cells) >= 4:
                if _SDI_SINGLE_RE.match(cells[_SDI_COL_4]):
                    sdi_value = cells[_SDI_COL_4]

            # Insert both the catalog (suffix-stripped) form and the
            # original form; overlay-merge so later tables only override
            # fields they actually carry.
            bare = _R_SUFFIX_RE.sub("", raw_name)
            for key in (bare, raw_name):
                inner = key[4:] if key.upper().startswith("LBL_") else key
                entry = result.setdefault(inner, {"octal": "", "sdi": "", "ssm": ""})
                entry["octal"] = octal_value
                if sdi_value:
                    entry["sdi"] = sdi_value
                if ssm_value:
                    entry["ssm"] = ssm_value
            # 8-col RDCU1 catalog (Table[0]): col 5 holds the bare signal
            # names carried by this LBL. Each is a separate mapping entry
            # so HLR text that references a signal by its bare name (no
            # ``LBL_`` prefix) still gets aliased to ``L<octal>_<signal>``
            # and inherits the row's SDI/SSM.
            if len(cells) >= _MIN_SIGNAL_COLS:
                for sig in _extract_signal_names(cells[_SIGNAL_COL_INDEX]):
                    sentry = result.setdefault(sig, {"octal": "", "sdi": "", "ssm": ""})
                    sentry["octal"] = octal_value
                    if sdi_value:
                        sentry["sdi"] = sdi_value
                    if ssm_value:
                        sentry["ssm"] = ssm_value
    return result


def preprocess_hlr_requirements(
    profile: ControllerProfile,
    hlr_out,
) -> int:
    """Rewrite HSCU HLR ``content`` to expose octal-form label aliases.

    Iterates over ``hlr_out.requirements`` and, for each requirement whose
    ``content`` mentions any mapped ``LBL_<NAME>`` (with optional ``_SSM``
    suffix), appends a per-requirement ``（亦称：L<octal>_<NAME>、...）``
    annotation listing ONLY the aliases for LBL tokens that actually appear
    in that requirement's content.

    Per-requirement scoping (vs. a single global suffix) prevents annotating
    a requirement with aliases for LBL names it does not reference, which
    would otherwise mislead the AI labeler into emitting label candidates
    that have no in-text support.

    Returns the number of requirements whose content was rewritten.
    """
    if not profile.hlr_preprocess.enabled:
        return 0

    mappings = profile.hlr_preprocess.extra_mappings
    cfg = profile.hlr_preprocess

    # Per-name meta lookup: {"octal": ..., "sdi": ..., "ssm": ...}.
    meta: dict[str, dict] = {}

    # Auto-extract from the HLR Word's LBL catalog tables if enabled. The
    # tables are auto-detected by _identify_label_tables() — no fixed
    # index required (handles HSCU documents where the LBL catalogs are
    # Table[0] 8-col RDCU1 inbound + Table[8] 4-col HSCU outbound, with
    # different column layouts). Acts as the base; explicit
    # extra_mappings extend/override below.
    if cfg.auto_parse_hlr_table_0 and hlr_out and getattr(hlr_out, "source_file", None):
        meta.update(
            _extract_label_meta(hlr_out.source_file)
        )

    # Apply explicit extra_mappings on top (octal only — they carry no
    # SDI/SSM info).
    for full_key, octal in mappings:
        octal_clean = octal.strip()
        if not octal_clean or "?" in octal_clean:
            continue
        bare_name = full_key[4:] if full_key.upper().startswith("LBL_") else full_key
        octal_value = (
            octal_clean[1:] if octal_clean.upper().startswith("L") else octal_clean
        )
        entry = meta.setdefault(bare_name, {"octal": "", "sdi": "", "ssm": ""})
        entry["octal"] = octal_value

    if not meta:
        return 0

    # Octal map with placeholder ("???") or empty values dropped.
    bare_octal: dict[str, str] = {
        name: m["octal"]
        for name, m in meta.items()
        if m.get("octal") and "?" not in m["octal"]
    }

    rewritten = 0
    for req in hlr_out.requirements:
        content = req.content
        if not content:
            continue
        # Two-pass token scan:
        #   1) ``LBL_<NAME>`` tokens (catalog name form)
        #   2) bare uppercase signal tokens that match a mapping entry
        #      directly (RDCU1 catalog col-5 signal names appearing in
        #      HLR body text without the ``LBL_`` prefix).
        lbl_tokens = _LBL_TOKEN.findall(content)
        # Only keep bare signal tokens that resolve to a mapping entry;
        # otherwise we'd noise up the suffix with arbitrary
        # uppercase-looking words.
        bare_signal_tokens = [
            t for t in _BARE_SIGNAL_TOKEN.findall(content)
            if t in bare_octal
        ]
        if not lbl_tokens and not bare_signal_tokens:
            continue
        # Only append aliases for tokens that, after stripping a trailing
        # ``_SSM`` (for the LBL path), resolve to a configured mapping.
        # Preserve declaration order so the suffix is stable across runs.
        # Each alias carries its catalog SDI/SSM when available:
        #   L145_DIS_00_SYS1（SDI=1，SSM类型=DIS）
        # The ``SDI=<n>`` form feeds the matcher's value-level SDI
        # dimension and Gate 2; ``SSM类型=`` is judge-facing context only.
        def _alias_for(name: str) -> str:
            m = meta.get(name, {})
            alias = f"L{_format_alias_octal(bare_octal[name])}_{name}"
            extras = []
            if m.get("sdi"):
                extras.append(f"SDI={m['sdi']}")
            if m.get("ssm"):
                extras.append(f"SSM类型={m['ssm']}")
            if extras:
                alias += "（" + "，".join(extras) + "）"
            return alias

        present_mapped: list[str] = []
        seen: set[str] = set()
        for token_name in lbl_tokens:
            bare = _bare_name(token_name)
            if bare in bare_octal and bare not in seen:
                seen.add(bare)
                present_mapped.append(_alias_for(bare))
        for sig in bare_signal_tokens:
            if sig not in seen:
                seen.add(sig)
                present_mapped.append(_alias_for(sig))
        if not present_mapped:
            continue
        suffix = "（亦称：" + "、".join(present_mapped) + "）"
        if suffix in content:
            continue
        req.content = content.rstrip() + " " + suffix
        rewritten += 1
    return rewritten