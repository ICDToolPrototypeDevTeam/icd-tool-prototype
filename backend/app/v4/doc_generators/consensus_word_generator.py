# -*- coding: utf-8 -*-
"""Word document generator: multi-model consensus review report with star ratings."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _set_cell_font(cell, text: str, bold: bool = False, size: int = 9, color=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _style_header_row(table, headers: list[str]):
    """Style the header row of a table."""
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_font(header_cells[i], h, bold=True, size=9)
        shading = header_cells[i]._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "D9E2F3",
            qn("w:val"): "clear",
        })
        shading.insert(0, shd)


def _star_str(n: int) -> str:
    """Render star rating as visual string (5-star system, ADR-004).

    Examples:
        5 → ★★★★★
        4 → ★★★★☆
        3 → ★★★☆☆
        2 → ★★☆☆☆
        1 → ★☆☆☆☆
        0 → ☆☆☆☆☆ (无匹配等不参与星评的 case)
    """
    n = max(0, min(5, int(n)))
    return "★" * n + "☆" * (5 - n)


def _map_consensus_label(stars: int, agreement: str) -> tuple[str, "RGBColor | None"]:
    """Map (stars, agreement) to (共识列文案, 颜色) — 与顶部"星级分布"表口径一致。

    3 档颜色（基于"采纳安全度"）：
    - 绿 (0x00, 0x80, 0x00)：5★ 完全共识 → 可直接采纳
    - 黄 (0xCC, 0x88, 0x00)：4★/3★ 有提醒但仍可采纳（4★ 完全共识·字段异议 / 3★ 多数共识）
    - 红 (0xCC, 0x00, 0x00)：2★/1★ 需关注或人工复核

    与顶部"星级分布"小节 (line 218-227) 的 5 档标签口径一致；1★ 3 个降级子类型
    按 agreement_level 区分（split → 三方分歧 / single_source → 仅单源 /
    no_consensus → 全部失效）。
    """
    GREEN = RGBColor(0x00, 0x80, 0x00)
    YELLOW = RGBColor(0xCC, 0x88, 0x00)
    RED = RGBColor(0xCC, 0x00, 0x00)
    if stars == 5:
        return ("完全共识", GREEN)
    if stars == 4:
        return ("完全共识·字段异议", YELLOW)
    if stars == 3:
        return ("多数共识", YELLOW)
    if stars == 2:
        return ("多数共识·关键异议", RED)
    if stars == 1:
        one_star_labels = {
            "split": ("三方分歧", RED),
            "single_source": ("仅单源", RED),
            "no_consensus": ("全部失效", RED),
        }
        return one_star_labels.get(agreement, ("降级", RED))
    return ("—", None)


def generate_consensus_report(
    consensus_path: Path,
    match_path: Path,
    output_path: Path,
) -> None:
    """Generate a multi-model consensus analysis Word document.

    Reads consensus_results.json + reverse_matches.json and produces
    a report with star ratings per HLR, agreement metrics, and per-model
    judgment details.

    Columns: 序号 | HLR ID | 信号类别 | 判定 | 匹配类型
             ICD Block | 分析摘要 | 共识 | 星级
    """
    consensus_data = json.loads(consensus_path.read_text(encoding="utf-8"))
    match_data = json.loads(match_path.read_text(encoding="utf-8"))

    consensus_results = consensus_data.get("results", [])
    consensus_summary = consensus_data.get("summary", {})
    match_results = match_data.get("results", [])

    # ── Build case_id → match info lookup ──
    # Match results in order: 已匹配 → 待确定 → 无匹配
    # Consensus results map to the first two groups (non-"无匹配")
    case_match_map: dict[str, dict] = {}
    match_idx = 0
    for mr in match_results:
        if mr.get("match_type") in ("已匹配", "待确定"):
            if match_idx < len(consensus_results):
                cid = consensus_results[match_idx]["case_id"]
                case_match_map[cid] = mr
                match_idx += 1

    # ── Merge consensus + match (judged entries) ──
    merged: list[dict] = []
    for cr in consensus_results:
        cid = cr["case_id"]
        mr = case_match_map.get(cid, {})
        # ADR-004 v3 fusion: inconsistent_attributes 为主字段（EoICD-HLR 事实差异，
        # 渲染到"不一致属性"列）；field_disagreements 仅入 JSON，不渲染
        merged.append({
            "case_id": cid,
            "hlr_id": mr.get("hlr_id", ""),
            "hlr_content": mr.get("hlr_content", ""),
            "matched_profile_keys": mr.get("matched_profile_keys", []),
            "agreement_level": cr.get("agreement_level", ""),
            "star_rating": cr.get("star_rating", 0),
            "final_coverage_status": cr.get("final_coverage_status", ""),
            "final_analysis": cr.get("final_analysis", ""),
            "confidence": cr.get("confidence", 0),
            "model_results": cr.get("model_results", {}),
            "field_disagreements": cr.get("field_disagreements", []),
            "inconsistent_attributes": cr.get("inconsistent_attributes", []),
            "is_no_match": False,
        })

    # ── Append "无匹配" entries (not judged, for full coverage) ──
    no_match_entries: list[dict] = []
    seq_after_judged = len(merged)
    for mr in match_results:
        if mr.get("match_type") == "无匹配":
            seq_after_judged += 1
            no_match_entries.append({
                "case_id": f"NOMATCH-{len(no_match_entries) + 1:04d}",
                "hlr_id": mr.get("hlr_id", ""),
                "hlr_content": mr.get("hlr_content", ""),
                "matched_profile_keys": [],
                "agreement_level": "",
                "star_rating": 0,
                "final_coverage_status": "无匹配",
                "final_analysis": mr.get("summary") or "匹配层未找到对应EoICD信号",
                "confidence": 0,
                "model_results": {},
                "inconsistent_attributes": [],
                "is_no_match": True,
            })
    merged_all = merged + no_match_entries

    # ── Build document ──
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    # Title
    title = doc.add_heading("EoICD 与 SWHLR 多模型差异分析报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"生成时间: {consensus_data.get('generated_at', 'N/A')[:19]}"
    )

    # ── Summary section ──
    doc.add_heading("一、差异分析总结", level=2)

    total_hlr = match_data.get("total_hlr", 0)
    match_stats = match_data.get("stats", {})
    cs = consensus_summary

    overview_parts = [
        f"对 {total_hlr} 条软件高层需求(SWHLR)进行多模型差异分析：",
        f"3 个裁判模型（DeepSeek / MiniMax / Qwen）并行独立判定，",
        f"1 个 Review Agent 对结果综合给出星级评价。",
    ]
    doc.add_paragraph("".join(overview_parts))

    # ── Summary tables ──
    doc.add_heading("判定分布", level=3)

    status_dist = cs.get("status_distribution", {})
    st = doc.add_table(rows=1, cols=2)
    st.style = "Table Grid"
    _style_header_row(st, ["判定结果", "数量"])

    # 数据驱动：遍历汇总键渲染，保证任何判定状态都不会从合计中静默消失
    status_colors = {
        "已覆盖": RGBColor(0x00, 0x80, 0x00),
        "不一致": RGBColor(0xCC, 0x33, 0x00),
        "待确认": RGBColor(0xCC, 0x55, 0x00),
    }
    ordered_statuses = ["已覆盖", "不一致", "待确认"]
    status_keys = ordered_statuses + sorted(
        k for k in status_dist if k not in ordered_statuses
    )
    judge_total = 0
    for label in status_keys:
        count = status_dist.get(label, 0)
        if count > 0:
            row = st.add_row()
            _set_cell_font(row.cells[0], label, bold=True, color=status_colors.get(label))
            _set_cell_font(row.cells[1], str(count))
            judge_total += count

    # 无匹配（来自匹配层统计）
    unmatched = match_stats.get("hlr_无匹配", 0)
    if unmatched > 0:
        row = st.add_row()
        _set_cell_font(row.cells[0], "无匹配", bold=True, color=RGBColor(0x99, 0x33, 0xCC))
        _set_cell_font(row.cells[1], str(unmatched))

    row = st.add_row()
    _set_cell_font(row.cells[0], "合计", bold=True)
    _set_cell_font(row.cells[1], str(judge_total + unmatched), bold=True)

    for row_obj in st.rows:
        row_obj.cells[0].width = Cm(6.0)
        row_obj.cells[1].width = Cm(3.0)

    # ── Consensus quality section ──
    doc.add_heading("星级分布", level=3)

    agreement_dist = cs.get("agreement_distribution", {})
    star_dist = cs.get("star_distribution", {})
    avg_stars = cs.get("average_star_rating", 0)

    # 5 星体系（ADR-004 v2）：5 档独立行；1★ 由 3 个降级子类型（split/single_source/
    # no_consensus）细分，保证降级产生的共识类型不丢计数。其它未在主表中的
    # 星级以「非常规」追加。
    star_levels = [
        ("5", "完全共识",
         "3 个模型判断完全一致，无字段级别分歧"),
        ("4", "完全共识·字段异议",
         "3 个模型判断一致，分析文本提及字段差异"),
        ("3", "多数共识",
         "多数模型判断一致，少数意见仅涉及辅助字段或模糊表达"),
        ("2", "多数共识·关键异议",
         "多数模型判断一致，少数意见涉及关键字段"),
    ]
    one_star_subs = [
        ("split", "三方分歧", "3 个模型各持不同意见"),
        ("single_source", "仅单源", "仅 1 个模型有效，其余调用失败"),
        ("no_consensus", "全部失效", "全部模型调用失败"),
    ]

    qt = doc.add_table(rows=1, cols=4)
    qt.style = "Table Grid"
    _style_header_row(qt, ["星级", "共识等级", "数量", "说明"])

    for star_key, label, desc in star_levels:
        count = star_dist.get(star_key, 0)
        row = qt.add_row()
        _set_cell_font(row.cells[0], _star_str(int(star_key)), bold=True, size=10,
                       color=RGBColor(0xCC, 0x88, 0x00))
        _set_cell_font(row.cells[1], label, bold=True)
        _set_cell_font(row.cells[2], f"{count} 条")
        _set_cell_font(row.cells[3], desc)

    # 1★ 子类型行：标签格式与主行一致（加粗、默认颜色）；首行星列放 ★☆☆☆☆ 后纵向合并
    one_star_first_row_idx: int | None = None
    one_star_last_row_idx: int | None = None
    for i, (agr_key, sub_label, sub_desc) in enumerate(one_star_subs):
        sub_count = agreement_dist.get(agr_key, 0)
        srow = qt.add_row()
        if i == 0:
            _set_cell_font(srow.cells[0], _star_str(1), bold=True, size=10,
                           color=RGBColor(0xCC, 0x88, 0x00))
            one_star_first_row_idx = len(qt.rows) - 1
        _set_cell_font(srow.cells[1], sub_label, bold=True)
        _set_cell_font(srow.cells[2], f"{sub_count} 条")
        _set_cell_font(srow.cells[3], sub_desc)
        one_star_last_row_idx = len(qt.rows) - 1

    for star_key in sorted(
        (k for k in star_dist if k not in {"1", "2", "3", "4", "5"}),
        key=lambda k: int(k) if k.isdigit() else 999,
    ):
        count = star_dist.get(star_key, 0)
        if count > 0:
            row = qt.add_row()
            _set_cell_font(row.cells[0], f"{star_key}★", bold=True, size=10,
                           color=RGBColor(0xCC, 0x88, 0x00))
            _set_cell_font(row.cells[1], "非常规", bold=True)
            _set_cell_font(row.cells[2], f"{count} 条")
            _set_cell_font(row.cells[3], "LLM 输出非常规星级，请人工复核")

    row = qt.add_row()
    _set_cell_font(row.cells[1], "平均星级", bold=True)
    _set_cell_font(row.cells[2], f"{avg_stars:.2f}", bold=True)

    for row_obj in qt.rows:
        row_obj.cells[0].width = Cm(2.5)
        row_obj.cells[1].width = Cm(3.0)
        row_obj.cells[2].width = Cm(2.0)
        row_obj.cells[3].width = Cm(8.0)

    # 1★ 3 个子类型行在星列合并为一个 ★☆☆☆☆ 单元格（纵向居中）
    if one_star_first_row_idx is not None and one_star_last_row_idx is not None:
        merged_cell = qt.cell(one_star_first_row_idx, 0).merge(qt.cell(one_star_last_row_idx, 0))
        merged_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in merged_cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)

    # ── Suggestions ──
    doc.add_heading("处置建议", level=3)
    suggestions = [
        f"{_star_str(5)} 星条目：完全共识，可直接采纳结论。",
        f"{_star_str(4)} 星条目：完全共识·字段异议，可采纳但建议核对字段细节。",
        f"{_star_str(3)} 星条目：多数共识，可采纳多数结论。",
        f"{_star_str(2)} 星条目：多数共识·关键异议，少数涉及关键字段，建议人工复核少数意见。",
        f"{_star_str(1)} 星条目（三方分歧）：三方各持不同意见，建议人工逐条复核并做出最终判断。",
        f"{_star_str(1)} 星条目（仅单源）：仅 1 个模型有效（其余模型调用失败），结论仅供参考，建议人工确认。",
        f"{_star_str(1)} 星条目（全部失效）：全部模型调用失败，AI 结论不可用，必须人工复核。",
        "无匹配 条目：匹配层未找到对应EoICD信号，需确认HLR是否属于ICD接口范畴。",
    ]
    for s in suggestions:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_page_break()

    # ── Detail tables grouped by coverage status ──
    doc.add_heading("二、分析明细", level=2)
    judged_count = len(merged)
    no_match_count = len(no_match_entries)
    doc.add_paragraph(f"共 {len(merged_all)} 条记录"
                      f"（{judged_count} 条已裁判 + {no_match_count} 条无匹配）")

    headers = ["序号", "SWHLR ID", "判定", "ICD Block",
               "不一致属性", "分析摘要", "共识", "星级"]

    status_labels = {
        "covered": "已覆盖",
        "inconsistent": "不一致",
        "needs_review": "待确认",
        "无匹配": "无匹配",
        "待确认": "待确认",
    }
    status_colors = {
        "已覆盖": RGBColor(0x00, 0x80, 0x00),
        "不一致": RGBColor(0xCC, 0x33, 0x00),
        "待确认": RGBColor(0xCC, 0x55, 0x00),
        "无匹配": RGBColor(0x99, 0x33, 0xCC),
    }

    # Section definitions: (group_key, heading_title, filter_predicate)
    sections = [
        ("covered", "已覆盖",
         lambda m: not m["is_no_match"] and m["final_coverage_status"] == "covered"),
        ("inconsistent", "不一致",
         lambda m: not m["is_no_match"] and m["final_coverage_status"] == "inconsistent"),
        ("needs_review", "待确认",
         lambda m: not m["is_no_match"] and m["final_coverage_status"] in ("needs_review", "待确认")),
        ("no_match", "无匹配",
         lambda m: m["is_no_match"]),
    ]

    seq = 0
    for _group_key, heading_title, pred in sections:
        group = [m for m in merged_all if pred(m)]
        if not group:
            continue

        doc.add_heading(f"{heading_title}（{len(group)} 条）", level=3)

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Set all tables to 100% page width for consistent sizing
        _tblW = table._tbl.tblPr.find(qn('w:tblW'))
        if _tblW is not None:
            _tblW.set(qn('w:w'), '5000')
            _tblW.set(qn('w:type'), 'pct')
        _style_header_row(table, headers)

        for m in group:
            seq += 1
            row = table.add_row()
            status = m["final_coverage_status"]
            agreement = m["agreement_level"]
            stars = m["star_rating"]
            is_nm = m["is_no_match"]

            status_cn = status_labels.get(status, status)

            blocks_str = ", ".join(m["matched_profile_keys"][:5]) if m["matched_profile_keys"] else "—"
            if len(m["matched_profile_keys"]) > 5:
                blocks_str += f" ... (+{len(m['matched_profile_keys']) - 5})"

            # 共识列标签按 (stars, agreement) 联合映射，与顶部"星级分布"表口径一致
            consensus_label, consensus_color = _map_consensus_label(stars, agreement)

            # Inconsistency attributes column
            attr_list = m.get("inconsistent_attributes", [])
            if attr_list:
                attr_str = " | ".join(a.get("attribute", "") for a in attr_list)
            else:
                attr_str = "—"

            _set_cell_font(row.cells[0], str(seq))
            _set_cell_font(row.cells[1], m["hlr_id"])
            _set_cell_font(row.cells[2], status_cn, bold=True, color=status_colors.get(status_cn))
            _set_cell_font(row.cells[3], blocks_str, size=7)
            _set_cell_font(row.cells[4], attr_str, size=7)
            _set_cell_font(row.cells[5], m["final_analysis"], size=8)
            if is_nm:
                _set_cell_font(row.cells[6], "—")
                _set_cell_font(row.cells[7], "—")
            else:
                _set_cell_font(row.cells[6], consensus_label, color=consensus_color)
                _set_cell_font(row.cells[7], _star_str(stars), bold=True, size=10,
                               color=RGBColor(0xCC, 0x88, 0x00))

        # 设置明细表列宽
        detail_col_widths = [1.19, 5.25, 1.5, 4.5, 2.0, 10.5, 1.75, 1.39]
        for row_obj in table.rows:
            for i, w in enumerate(detail_col_widths):
                row_obj.cells[i].width = Cm(w)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  Consensus report: {output_path}")
