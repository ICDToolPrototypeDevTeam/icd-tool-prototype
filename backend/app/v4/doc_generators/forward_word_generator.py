# -*- coding: utf-8 -*-
"""Forward completeness Word report (Stage C8).

Six sections:
  一、分析说明与判定规则  二、分析结果概览  三、未覆盖清单
  四、待确认清单          五、输入异常对象清单  六、不支持清单

「已覆盖」只在结果概览中统计；后四个清单严格按最终状态互斥分类。每个清单统一
六列：EoICD ID / 协议 / 信号族 / 设备 / 覆盖状态 / 原因。原因用自然语言（缺失
候选只显示数量），完整 ID 仅保留在 Excel / JSON。
"""

from __future__ import annotations

from pathlib import Path

from app.v4.doc_generators.forward_excel_generator import _coverage_detail_row
from app.v4.models import ForwardBlocksOutput, ForwardCoverageOutput

_HEADERS = ["EoICD ID", "协议", "信号族", "设备", "覆盖状态", "原因"]
_COL_WIDTHS_CM = [5.0, 1.8, 5.0, 3.3, 3.5, 9.0]


def _section_of(row: dict) -> str:
    """Map a row to its single section (mutually exclusive by final status)."""
    if row["analysis_status"] == "unsupported":
        return "unsupported"
    if row["analysis_status"] == "input_error":
        return "input_error"
    if row["coverage_status"] == "uncovered":
        return "uncovered"
    if row["coverage_status"] in ("possible", "parent_referenced"):
        return "pending"
    return "covered"


def _global_missing_hlrs(blocks: ForwardBlocksOutput) -> set[str]:
    """去重后的、追溯表引用但未出现在上传 HLR 文档中的候选 HLR 总数。"""
    missing: set[str] = set()
    for b in blocks.blocks:
        if b.trace:
            missing.update(b.trace.missing_hlr_ids)
    return missing


def generate_forward_word(
    coverage: ForwardCoverageOutput,
    blocks: ForwardBlocksOutput,
    output_path: Path,
) -> None:
    """Generate the forward completeness Word report."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    block_map = {b.business_object_id: b for b in blocks.blocks}
    rows = [
        _coverage_detail_row(block_map[r.business_object_id], r)
        for r in coverage.results
        if r.business_object_id in block_map
    ]

    covered = [r for r in rows if _section_of(r) == "covered"]
    uncovered = [r for r in rows if _section_of(r) == "uncovered"]
    pending = [r for r in rows if _section_of(r) == "pending"]
    input_errors = [r for r in rows if _section_of(r) == "input_error"]
    unsupported = [r for r in rows if _section_of(r) == "unsupported"]

    global_missing = len(_global_missing_hlrs(blocks))

    black = RGBColor(0x00, 0x00, 0x00)
    red = RGBColor(0xCC, 0x33, 0x00)
    orange = RGBColor(0xCC, 0x55, 0x00)

    # ── Font helpers (reliable CJK fallback across Word / LibreOffice) ──────
    def apply_font(run, size=9, bold=False, color=None):
        run.font.name = "微软雅黑"
        run.font.size = Pt(size)
        run.bold = bold
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "微软雅黑")
        rfonts.set(qn("w:hAnsi"), "微软雅黑")
        rfonts.set(qn("w:eastAsia"), "微软雅黑")
        rfonts.set(qn("w:cs"), "微软雅黑")
        if color is not None:
            run.font.color.rgb = color

    def add_heading(text, level=2, page_break=False):
        h = doc.add_heading("", level=level)
        run = h.add_run(text)
        apply_font(run, size=(16 if level == 1 else 12), bold=True, color=black)
        if page_break:
            h.paragraph_format.page_break_before = True
        return h

    def add_para(text="", bold=False):
        p = doc.add_paragraph()
        if text:
            apply_font(p.add_run(text), size=9, bold=bold)
        return p

    def add_labeled_para(label, text):
        p = doc.add_paragraph()
        apply_font(p.add_run(label + "："), size=9, bold=True)
        apply_font(p.add_run(text), size=9, bold=False)
        return p

    def set_font(cell, text, bold=False, size=9, color=None):
        cell.text = ""
        p = cell.paragraphs[0]
        apply_font(p.add_run(str(text)), size=size, bold=bold, color=color)

    def style_header(table, headers):
        for i, h in enumerate(headers):
            set_font(table.rows[0].cells[i], h, bold=True, size=9)
            tc = table.rows[0].cells[i]._element.get_or_add_tcPr()
            shd = tc.makeelement(qn("w:shd"), {qn("w:fill"): "D9E2F3", qn("w:val"): "clear"})
            tc.insert(0, shd)

    def set_col_widths(table, widths_cm):
        table.autofit = False
        for row in table.rows:
            for i, w in enumerate(widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    def fill_table(table, headers, data, color=None):
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        style_header(table, headers)
        for r in data:
            row = table.add_row()
            set_font(row.cells[0], r["business_object_id"])
            set_font(row.cells[1], r["protocol"])
            set_font(row.cells[2], r["signal_family"] or r["signal"])
            set_font(row.cells[3], r["device"])
            set_font(row.cells[4], r["coverage_label"], bold=True, color=color)
            set_font(row.cells[5], r["reason"] or "—")
        set_col_widths(table, _COL_WIDTHS_CM)

    def add_list_section(title, description, data, color=None, page_break=False):
        add_heading(title, level=2, page_break=page_break)
        add_para(description)
        if data:
            t = doc.add_table(rows=1, cols=len(_HEADERS))
            t.style = "Table Grid"
            fill_table(t, _HEADERS, data, color=color)
        else:
            add_para("（无）")

    # ── Document setup ──────────────────────────────────────────────────────
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(9)
    normal_rpr = normal.element.get_or_add_rPr()
    normal_rfonts = normal_rpr.get_or_add_rFonts()
    normal_rfonts.set(qn("w:eastAsia"), "微软雅黑")

    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    title = add_heading("EoICD 正向完整性分析报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        f"分析模式: {coverage.analysis_mode or 'N/A'}    "
        f"生成时间: {coverage.generated_at[:19]}"
    )
    add_para(
        f"AI 调用统计: HLR 标签调用 {coverage.hlr_label_calls} 次 / "
        f"正向三态复核调用 {coverage.ai_review_calls} 次"
    )

    # ── 一、分析说明与判定规则 ──────────────────────────────────────────────
    add_heading("一、分析说明与判定规则")
    add_para(
        "本报告从 EoICD 源文件出发，检查每个业务对象（业务信号/字段）在软件高层需求（HLR）"
        "正文中是否存在对应描述，用于识别「漏写」。判定按确定性规则与 AI 复核两层进行："
        "先依据对象的 Label 号、信号名、SDI、bit 等身份信息与 HLR 描述做确定性比对，"
        "确定性规则无法定论时再交由 AI 复核；当追溯表引用的候选 HLR 未出现在上传的 HLR 文档时，"
        "优先保守处理为「待确认」，避免误报漏写。"
    )
    for label, desc in [
        ("已覆盖", "HLR 正文中存在该 EoICD 业务对象的对应描述。"),
        ("待确认", "存在候选 HLR 但无法确定是否描述了该对象，需人工审查。"),
        ("未覆盖", "HLR 正文中未找到该对象的对应描述，疑似漏写。"),
        ("输入异常对象", "追溯表引用的候选 HLR 全部缺失于上传的 HLR 文档，无法分析。"),
        ("不支持", "该对象的协议类型（如原生 A664）暂不支持分析。"),
    ]:
        add_labeled_para(label, desc)
    add_para(
        "清单划分说明：已覆盖只在结果概览中统计，不单独展开；未覆盖、待确认、输入异常对象、"
        "不支持四个清单严格按最终状态分类，一个对象只进入其中一个清单。"
    )

    # ── 二、分析结果概览 ────────────────────────────────────────────────────
    add_heading("二、分析结果概览")
    total = len(rows)
    ot = doc.add_table(rows=1, cols=3)
    ot.style = "Table Grid"
    style_header(ot, ["覆盖状态", "数量", "占比"])
    for label, count in [
        ("已覆盖", len(covered)),
        ("待确认", len(pending)),
        ("未覆盖", len(uncovered)),
        ("输入异常对象", len(input_errors)),
        ("不支持", len(unsupported)),
    ]:
        row = ot.add_row()
        pct = f"{count / total * 100:.1f}%" if total else "0%"
        set_font(row.cells[0], label, bold=True)
        set_font(row.cells[1], str(count))
        set_font(row.cells[2], pct)
    row = ot.add_row()
    set_font(row.cells[0], "合计", bold=True)
    set_font(row.cells[1], str(total), bold=True)
    set_font(row.cells[2], "100%")

    if global_missing:
        add_labeled_para(
            "全局输入警告",
            f"追溯表引用但未出现在上传 HLR 文档中的候选 HLR 共 {global_missing} 条（去重后），"
            f"这些缺失可能影响部分对象的最终覆盖结论。详细 ID 见 Excel「缺失HLR明细」表。",
        )

    # ── 三~六、清单（三 开始换页，避免手动分页符与自动分页叠加产生空白页）──
    add_list_section(
        "三、未覆盖清单",
        f"共 {len(uncovered)} 个 EoICD 业务对象未在 HLR 中找到对应描述。",
        uncovered, color=red, page_break=True,
    )
    add_list_section(
        "四、待确认清单",
        f"共 {len(pending)} 个对象存在候选但无法确定，需人工审查。",
        pending, color=orange,
    )
    add_list_section(
        "五、输入异常对象清单",
        f"共 {len(input_errors)} 个对象因追溯表引用的候选 HLR 全部缺失于上传的 HLR 文档，无法分析。",
        input_errors,
    )
    add_list_section(
        "六、不支持清单",
        f"共 {len(unsupported)} 个对象因协议类型暂不支持分析。",
        unsupported,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  Forward Word: {output_path}")
