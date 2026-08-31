# -*- coding: utf-8 -*-
"""Word document generator: DeepSeek-only consistency analysis report."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _set_cell_font(cell, text: str, bold: bool = False, size: int = 9, color=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _style_header_row(table, headers: list[str]):
    """Style the header row of a table."""
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_font(header_cells[i], h, bold=True, size=9)
        shading = header_cells[i]._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "D9E2F3",
            qn("w:val"): "clear",
        })
        shading.insert(0, shd)


_STATUS_LABELS = {
    "covered": "已覆盖",
    "inconsistent": "不一致",
    "needs_review": "需确认",
    "无匹配": "无匹配",
}

_STATUS_DESCRIPTIONS = [
    ("已覆盖", "SWHLR 中明确声明了与 ICD 定义一致的技术内容，ICD 接口要求在 SWHLR 中正确落实。"),
    ("不一致", "SWHLR 中的声明与 ICD 定义存在矛盾，如数据类型、位含义、方向、数值范围等不一致，需要修正。"),
    ("需确认", "SWHLR 与 ICD 之间存在模糊或不确定的对应关系，或匹配的 ICD Block 与 SWHLR 内容不直接相关，需人工审查。"),
    ("无匹配", "匹配阶段未在 EoICD 中找到对应的 ICD 信号定义，可能为 SWHLR 超出 ICD 范围的内容，需人工确认。"),
]

_MODEL_DISPLAY = {
    "deepseek": "DeepSeek",
    "minimax": "MiniMax",
    "qwen": "Qwen",
}

_STATUS_COLORS = {
    "covered": RGBColor(0x00, 0x80, 0x00),
    "inconsistent": RGBColor(0xCC, 0x33, 0x00),
    "needs_review": RGBColor(0xCC, 0x55, 0x00),
    "无匹配": RGBColor(0x99, 0x33, 0xCC),
}


def _load_hlr_metadata(report_dir: Path) -> tuple[dict[str, dict], list[dict]]:
    """Load reverse_matches.json and build lookup + ordered AI-judged HLR list."""
    match_path = report_dir / "reverse_matches.json"
    hlr_lookup: dict[str, dict] = {}
    ai_matched: list[dict] = []

    if not match_path.exists():
        return hlr_lookup, ai_matched

    match_data = json.loads(match_path.read_text(encoding="utf-8"))
    for m in match_data.get("results", []):
        hlr_lookup[m["hlr_id"]] = m
        if m.get("match_type") in ("已匹配", "待确定"):
            ai_matched.append(m)

    return hlr_lookup, ai_matched


def generate_consistency_report(
    reverse_report_path: Path,
    output_path: Path,
    model: str = "deepseek",
) -> None:
    """Generate single-model consistency analysis Word report.

    Reads reverse_report.json (multi-model results) and reverse_matches.json
    (HLR metadata) from the same directory. Extracts only the specified model's judgments.

    Args:
        reverse_report_path: Path to reverse_report.json
        output_path: Where to write the .docx file
        model: Which model's results to extract ("deepseek", "minimax", "qwen")
    """
    data = json.loads(reverse_report_path.read_text(encoding="utf-8"))
    summary = data.get("summary", {})
    results = data.get("results", [])

    model_display = _MODEL_DISPLAY.get(model, model)

    hlr_lookup, ai_matched_hlrs = _load_hlr_metadata(reverse_report_path.parent)

    # ── Build DeepSeek-only detailed results ──
    detailed: list[dict] = []
    status_counts: dict[str, int] = {}

    for r in results:
        case_id = r.get("case_id", "")

        if case_id.startswith("REV-"):
            idx = int(case_id.split("-")[1]) - 1
            hlr_info = ai_matched_hlrs[idx] if idx < len(ai_matched_hlrs) else {}

            ds = r.get("match_evidence", {}).get("model_results", {}).get(model, {})
            status = ds.get("coverage_status", "")
            entry = {
                "hlr_id": hlr_info.get("hlr_id", ""),
                "hlr_content": hlr_info.get("hlr_content", ""),
                "signal_category": hlr_info.get("signal_category", ""),
                "icd_blocks": hlr_info.get("matched_profile_keys", []),
                "status": status,
                "analysis": ds.get("analysis", ""),
                "confidence": ds.get("confidence", 0),
            }
        elif case_id.startswith("NOMATCH-"):
            status = "无匹配"
            entry = {
                "hlr_id": r.get("hlr_id", ""),
                "hlr_content": r.get("hlr_content", ""),
                "signal_category": "",
                "icd_blocks": [],
                "status": status,
                "analysis": r.get("analysis", ""),
                "confidence": 0,
            }
        else:
            continue

        detailed.append(entry)
        status_counts[status] = status_counts.get(status, 0) + 1

    # ── Build Word document ──
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    title = doc.add_heading("EoICD 与 SWHLR 单模型差异分析报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"生成时间: {data.get('generated_at', 'N/A')[:19]}\n"
        f"裁判模型: {model_display} (单模型)"
    )

    # ═══════════════════════════════════════════
    # Section 1: 判定结果概览表
    # ═══════════════════════════════════════════
    doc.add_heading("一、判定结果概览", level=2)

    doc.add_heading("判定分类说明", level=3)
    for label, desc in _STATUS_DESCRIPTIONS:
        p = doc.add_paragraph()
        p.add_run(f"{label}：").bold = True
        p.add_run(desc)

    overview_headers = ["判定结果", "数量", "占比"]
    ot = doc.add_table(rows=1, cols=3)
    ot.style = "Table Grid"
    _style_header_row(ot, overview_headers)

    total = len(detailed)
    for status_key in ("covered", "inconsistent", "needs_review", "无匹配"):
        count = status_counts.get(status_key, 0)
        pct = f"{count / total * 100:.1f}%" if total > 0 else "0%"
        color = _STATUS_COLORS.get(status_key)
        row = ot.add_row()
        _set_cell_font(row.cells[0], _STATUS_LABELS.get(status_key, status_key),
                       bold=True, color=color)
        _set_cell_font(row.cells[1], str(count))
        _set_cell_font(row.cells[2], pct)

    row = ot.add_row()
    _set_cell_font(row.cells[0], "合计", bold=True)
    _set_cell_font(row.cells[1], str(total), bold=True)
    _set_cell_font(row.cells[2], "100%")

    for row_obj in ot.rows:
        row_obj.cells[0].width = Cm(6.0)
        row_obj.cells[1].width = Cm(2.0)
        row_obj.cells[2].width = Cm(2.0)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # Section 2: 分析明细表
    # ═══════════════════════════════════════════
    doc.add_heading("二、分析明细", level=2)
    doc.add_paragraph(f"共 {total} 条记录（仅展示 {model_display} 裁判结果）")

    detail_headers = ["序号", "SWHLR ID", "判定结果", "ICD Block", "分析摘要", "置信度"]
    dt = doc.add_table(rows=1, cols=len(detail_headers))
    dt.style = "Table Grid"
    dt.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_header_row(dt, detail_headers)

    for i, entry in enumerate(detailed, 1):
        row = dt.add_row()
        status = entry["status"]
        color = _STATUS_COLORS.get(status)

        blocks = entry["icd_blocks"]
        block_str = ", ".join(blocks) if blocks else "—"

        conf_str = f"{entry['confidence']:.2f}" if entry["confidence"] > 0 else "—"

        _set_cell_font(row.cells[0], str(i))
        _set_cell_font(row.cells[1], entry["hlr_id"])
        _set_cell_font(row.cells[2], _STATUS_LABELS.get(status, status),
                       bold=True, color=color)
        _set_cell_font(row.cells[3], block_str, size=7)
        _set_cell_font(row.cells[4], entry["analysis"], size=8)
        _set_cell_font(row.cells[5], conf_str)

    col_widths = [0.8, 5.0, 1.9, 4.5, 12.0, 1.2]
    for row_obj in dt.rows:
        for i, w in enumerate(col_widths):
            row_obj.cells[i].width = Cm(w)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  Consistency report: {output_path}")
